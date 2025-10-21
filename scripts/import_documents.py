#!/usr/bin/env python3
"""
Script to import documents (PDF, DOC, DOCX) to the vector database.
Usage: python scripts/import_documents.py <directory>
"""

import os
import sys
import asyncio
from pathlib import Path
from datetime import datetime

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

import fitz  # PyMuPDF
from docx import Document
import docx2txt

from src.discord_bot.rag.vector_store import vector_store


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using PyMuPDF"""
    try:
        doc = fitz.open(file_path)
        text = ""
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            if isinstance(page_text, str) and page_text.strip():
                text += f"\n--- Página {page_num + 1} ---\n{page_text}"
        doc.close()
        return text.strip()
    except Exception as e:
        print(f"❌ Error reading PDF {file_path}: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        # Try with python-docx first
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Try to extract tables too
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)

        text = "\n".join(paragraphs)

        if tables_text:
            text += "\n\n--- Tables ---\n" + "\n".join(tables_text)

        # If nothing extracted, try with docx2txt
        if not text.strip():
            text = docx2txt.process(file_path)

        return text.strip()
    except Exception as e:
        print(f"❌ Error reading DOCX {file_path}: {e}")
        return ""


def extract_text_from_doc(file_path: str) -> str:
    """Extract text from DOC file (old format)"""
    try:
        # docx2txt can process .doc in some cases
        text = docx2txt.process(file_path)
        return text.strip()
    except Exception as e:
        print(f"⚠️  Warning: Could not read DOC {file_path}: {e}")
        print(f"   Old .doc files may need to be converted to .docx first")
        return ""


async def process_file(file_path: Path, collection_name: str = "rag_collection") -> bool:
    """Process a file and add to vector database"""

    file_extension = file_path.suffix.lower()

    print(f"\n📄 Processing: {file_path.name}")

    # Extract text based on file type
    if file_extension == ".pdf":
        text = extract_text_from_pdf(str(file_path))
    elif file_extension == ".docx":
        text = extract_text_from_docx(str(file_path))
    elif file_extension == ".doc":
        text = extract_text_from_doc(str(file_path))
    else:
        print(f"⚠️  Unsupported file type: {file_extension}")
        return False

    if not text:
        print(f"⚠️  No text extracted from {file_path.name}")
        return False

    # Information about extracted text
    words = len(text.split())
    chars = len(text)
    print(f"   📊 Text extracted: {words} words, {chars} characters")

    # Add to vector database
    try:
        metadata = {
            "filename": file_path.name,
            "filepath": str(file_path.absolute()),
            "filetype": file_extension,
            "imported_at": datetime.now().isoformat(),
            "word_count": words,
            "char_count": chars
        }

        await vector_store.add_document(
            text=text,
            metadata=metadata,
            collection_name=collection_name
        )

        print(f"   ✅ Added to vector database!")
        return True

    except Exception as e:
        print(f"   ❌ Error adding to vector database: {e}")
        return False


async def import_documents_from_directory(directory: str, collection_name: str = "rag_collection"):
    """Import all supported documents from a directory"""

    directory_path = Path(directory)

    if not directory_path.exists():
        print(f"❌ Error: Directory '{directory}' not found")
        return

    if not directory_path.is_dir():
        print(f"❌ Error: '{directory}' is not a directory")
        return

    print(f"\n🔍 Searching for documents in: {directory_path.absolute()}")

    # List of supported extensions
    supported_extensions = {".pdf", ".docx", ".doc"}

    # Find all supported files
    files = []
    for ext in supported_extensions:
        files.extend(directory_path.glob(f"*{ext}"))

    if not files:
        print(f"\n⚠️  No documents found (.pdf, .docx, .doc)")
        return

    print(f"📚 Found {len(files)} document(s)")
    print("=" * 60)

    # Process each file
    success_count = 0
    failed_count = 0

    for file_path in sorted(files):
        success = await process_file(file_path, collection_name)
        if success:
            success_count += 1
        else:
            failed_count += 1

    # Summary
    print("\n" + "=" * 60)
    print(f"\n📊 Import summary:")
    print(f"   ✅ Success: {success_count} document(s)")
    print(f"   ❌ Failed: {failed_count} document(s)")
    print(f"   📦 Total: {len(files)} file(s) processed")

    # Show vector database statistics
    try:
        count = vector_store.count_documents(collection_name)
        print(f"\n💾 Total documents in vector database: {count}")
    except Exception as e:
        print(f"\n⚠️  Error getting statistics: {e}")


async def main():
    """Main function"""

    print("=" * 60)
    print("📚 DOCUMENT IMPORTER - Discord Bot RAG")
    print("=" * 60)

    # Check arguments
    if len(sys.argv) < 2:
        print("\n❌ Usage: python scripts/import_documents.py <directory>")
        print("\nExample:")
        print("  python scripts/import_documents.py ./documents")
        print("  python scripts/import_documents.py /home/user/my_docs")
        print("\nSupported formats: PDF, DOCX, DOC")
        sys.exit(1)

    directory = sys.argv[1]

    # Initialize vector database
    print("\n🔧 Initializing vector database...")
    vector_store.init_vector_db()

    # Import documents
    await import_documents_from_directory(directory)

    print("\n✅ Import completed!")
    print("=" * 60)


if __name__ == "__main__":
    asyncio.run(main())
