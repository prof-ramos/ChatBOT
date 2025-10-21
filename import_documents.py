#!/usr/bin/env python3
"""
Script para importar documentos (PDF, DOC, DOCX) para o banco vetorial do bot Discord.
Uso: python import_documents.py <diretório>
"""

import os
import sys
import asyncio
from pathlib import Path
from datetime import datetime
import fitz  # PyMuPDF
from docx import Document
import docx2txt

# Importa as funções do bot
import vector_db


def extract_text_from_pdf(file_path: str) -> str:
    """Extrai texto de um arquivo PDF usando PyMuPDF"""
    try:
        doc = fitz.open(file_path)
        text = ""
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            if isinstance(page_text, str) and page_text.strip():
                text += f"\n--- Página {page_num + 1} ---\n{page_text}"
        doc.close()
        return text.strip()
    except Exception as e:
        print(f"❌ Erro ao ler PDF {file_path}: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extrai texto de um arquivo DOCX usando python-docx e docx2txt"""
    try:
        # Tenta com python-docx primeiro
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Tenta extrair tabelas também
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)
        
        text = "\n".join(paragraphs)
        
        if tables_text:
            text += "\n\n--- Tabelas ---\n" + "\n".join(tables_text)
        
        # Se não conseguiu extrair nada, tenta com docx2txt
        if not text.strip():
            text = docx2txt.process(file_path)
        
        return text.strip()
    except Exception as e:
        print(f"❌ Erro ao ler DOCX {file_path}: {e}")
        return ""


def extract_text_from_doc(file_path: str) -> str:
    """Extrai texto de um arquivo DOC (formato antigo)"""
    try:
        # docx2txt pode processar .doc em alguns casos
        text = docx2txt.process(file_path)
        return text.strip()
    except Exception as e:
        print(f"⚠️  Aviso: Não foi possível ler DOC {file_path}: {e}")
        print(f"   Arquivos .doc antigos podem precisar ser convertidos para .docx primeiro")
        return ""


async def process_file(file_path: Path, collection_name: str = "rag_collection") -> bool:
    """Processa um arquivo e adiciona ao banco vetorial"""
    
    file_extension = file_path.suffix.lower()
    
    print(f"\n📄 Processando: {file_path.name}")
    
    # Extrai o texto baseado no tipo de arquivo
    if file_extension == ".pdf":
        text = extract_text_from_pdf(str(file_path))
    elif file_extension == ".docx":
        text = extract_text_from_docx(str(file_path))
    elif file_extension == ".doc":
        text = extract_text_from_doc(str(file_path))
    else:
        print(f"⚠️  Tipo de arquivo não suportado: {file_extension}")
        return False
    
    if not text:
        print(f"⚠️  Nenhum texto extraído de {file_path.name}")
        return False
    
    # Informações sobre o texto extraído
    words = len(text.split())
    chars = len(text)
    print(f"   📊 Texto extraído: {words} palavras, {chars} caracteres")
    
    # Adiciona ao banco vetorial
    try:
        metadata = {
            "filename": file_path.name,
            "filepath": str(file_path.absolute()),
            "filetype": file_extension,
            "imported_at": datetime.now().isoformat(),
            "word_count": words,
            "char_count": chars
        }
        
        await vector_db.add_document(
            text=text,
            metadata=metadata,
            collection_name=collection_name
        )
        
        print(f"   ✅ Adicionado ao banco vetorial!")
        return True
        
    except Exception as e:
        print(f"   ❌ Erro ao adicionar ao banco vetorial: {e}")
        return False


async def import_documents_from_directory(directory: str, collection_name: str = "rag_collection"):
    """Importa todos os documentos suportados de um diretório"""
    
    directory_path = Path(directory)
    
    if not directory_path.exists():
        print(f"❌ Erro: Diretório '{directory}' não encontrado")
        return
    
    if not directory_path.is_dir():
        print(f"❌ Erro: '{directory}' não é um diretório")
        return
    
    print(f"\n🔍 Procurando documentos em: {directory_path.absolute()}")
    
    # Lista de extensões suportadas
    supported_extensions = {".pdf", ".docx", ".doc"}
    
    # Encontra todos os arquivos suportados
    files = []
    for ext in supported_extensions:
        files.extend(directory_path.glob(f"*{ext}"))
    
    if not files:
        print(f"\n⚠️  Nenhum documento encontrado (.pdf, .docx, .doc)")
        return
    
    print(f"📚 Encontrados {len(files)} documento(s)")
    print("=" * 60)
    
    # Processa cada arquivo
    success_count = 0
    failed_count = 0
    
    for file_path in sorted(files):
        success = await process_file(file_path, collection_name)
        if success:
            success_count += 1
        else:
            failed_count += 1
    
    # Resumo
    print("\n" + "=" * 60)
    print(f"\n📊 Resumo da importação:")
    print(f"   ✅ Sucesso: {success_count} documento(s)")
    print(f"   ❌ Falhas: {failed_count} documento(s)")
    print(f"   📦 Total: {len(files)} arquivo(s) processados")
    
    # Mostra estatísticas do banco vetorial
    try:
        count = vector_db.count_documents(collection_name)
        print(f"\n💾 Total de documentos no banco vetorial: {count}")
    except Exception as e:
        print(f"\n⚠️  Erro ao obter estatísticas: {e}")


async def main():
    """Função principal"""
    
    print("=" * 60)
    print("📚 IMPORTADOR DE DOCUMENTOS - Bot Discord RAG")
    print("=" * 60)
    
    # Verifica argumentos
    if len(sys.argv) < 2:
        print("\n❌ Uso: python import_documents.py <diretório>")
        print("\nExemplo:")
        print("  python import_documents.py ./documentos")
        print("  python import_documents.py /home/user/meus_docs")
        print("\nFormatos suportados: PDF, DOCX, DOC")
        sys.exit(1)
    
    directory = sys.argv[1]
    
    # Inicializa o banco vetorial
    print("\n🔧 Inicializando banco vetorial...")
    vector_db.init_vector_db()
    
    # Importa os documentos
    await import_documents_from_directory(directory)
    
    print("\n✅ Importação concluída!")
    print("=" * 60)


if __name__ == "__main__":
    asyncio.run(main())
